from docx import Document
from docx.shared import Inches,Pt
import openpyxl
 
# Define variable to load the dataframe
dataframe = openpyxl.load_workbook("applicant-list.xlsx")
 
# Define variable to read sheet
dataframe1 = dataframe.active
 
# Iterate the loop to read the cell values
for row in range(1, dataframe1.max_row):

    dict1 = {
        'c1' : '',
        'c2' : '',
        'c3' : '',
        'c4' : '',
        'c5' : '',
        'c6' : '',
        'c7' : '',
        'c8' : '',
        'c9' : '',
        'c10' : '',
        'c11' : '',
        'c12' : '',
        'c13' : '',
        'c14' : '',
        'c15' : '',
        'c16' : '',
        'c17' : '',
        'c18' : '',
        'c19' : '',
        'c20' : '',
        'c21' : '',
        'c22' : '',
        'c23' : '',
        'c24' : '',
        'c25' : '',
        'c26' : '',
        'c27' : '',
        'c28' : '',
        'c29' : '',
        'c30' : '',
        'c31' : '',
        'c32' : '',
        'c33' : '',
        'c34' : '',
        'c35' : '',
        'c36' : '',
        'c37' : '',
        'c38' : '',
        'c39' : '',
        'c40' : ''
    }

    dict2 = {
        1 : 'c1',
        2 : 'c2',
        3 : 'c3',
        4 : 'c4',
        5 : 'c5',
        6 : 'c6',
        7 : 'c7',
        8 : 'c8',
        9 : 'c9',
        10 : 'c10',
        11 : 'c11',
        12 : 'c12',
        13 : 'c13',
        14 : 'c14',
        15 : 'c15',
        16 : 'c16',
        17 : 'c17',
        18 : 'c18',
        19 : 'c19',
        20 : 'c20',
        21 : 'c21',
        22 : 'c22',
        23 : 'c23',
        24 : 'c24',
        25 : 'c25',
        26 : 'c26',
        27 : 'c27',
        28 : 'c28',
        29 : 'c29',
        30 : 'c30',
        31 : 'c31',
        32 : 'c32',
        33 : 'c33',
        34 : 'c34',
        35 : 'c35',
        36 : 'c36',
        37 : 'c37',
        38 : 'c38',
        39 : 'c39',
        40 : 'c40'
    }

    count = 0
    for col in dataframe1.iter_cols(1, 40):
        count+=1
        currentColValue = col[row].value
        dict1[dict2[count]] = currentColValue

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    p = document.add_paragraph('')
    p.add_run('2023 Professional Fellowship Program (Fall)').bold = True

    p = document.add_paragraph('')
    p.add_run('Name: ').bold = True
    p.add_run(dict1['c1'])

    p = document.add_paragraph('')
    p.add_run('Gender: ').bold = True
    p.add_run(dict1['c2'])

    p = document.add_paragraph('')
    p.add_run('Date of birth: ').bold = True
    p.add_run(dict1['c3'])

    p = document.add_paragraph('')
    p.add_run('Religion: ').bold = True
    p.add_run(dict1['c4'])

    p = document.add_paragraph('')
    p.add_run('Marital status: ').bold = True
    p.add_run(dict1['c5'])

    p = document.add_paragraph('')
    p.add_run('List the names, relationship and contact details of ALL people living in your household:(Eg: Mr. Robert Carr, Father, Cell: 0123456789, Email: example@mail.com)').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c6'])

    p = document.add_paragraph('')
    p.add_run('Do you have your family’s permission to go on this program should you get selected? ').bold = True
    p.add_run(dict1['c7'])

    p = document.add_paragraph('')
    p.add_run('Do you have a passport in hand that is valid six months after the program implementation deadline? ').bold = True
    p.add_run(dict1['c8'])

    p = document.add_paragraph('')
    p.add_run('City of Birth: ').bold = True
    p.add_run(dict1['c9'])

    p = document.add_paragraph('')
    p.add_run('Primary contact information - ').bold = True
    p.add_run(dict1['c10'])

    p = document.add_paragraph('')
    p.add_run('Mobile Phone: ').bold = True
    p.add_run(dict1['c11'])

    p = document.add_paragraph('')
    p.add_run('Passport Number: ').bold = True
    p.add_run(dict1['c12'])

    p = document.add_paragraph('')
    p.add_run('Passport issuing authority: ').bold = True
    p.add_run(dict1['c13'])

    p = document.add_paragraph('')
    p.add_run('Expiry date: ').bold = True
    p.add_run(dict1['c14'])

    p = document.add_paragraph('')
    p.add_run('Current address: ').bold = True
    p.add_run(dict1['c15'])

    p = document.add_paragraph('')
    p.add_run('The U.S. government does not discriminate against applicants because of race, color, religion, sex, age, national origin, disability or any other protected characteristic as established by U.S. law. If selected to participate in the program, would you require any special assistance or accommodation due to a disability or special need?').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c16'])

    p = document.add_paragraph('')
    p.add_run('Current employer/Organization where you currently work: ').bold = True
    p.add_run(dict1['c17'])

    p = document.add_paragraph('')
    p.add_run('Category of work: ').bold = True
    p.add_run(dict1['c18'])

    p = document.add_paragraph('')
    p.add_run('How many years have you been at this organization? ').bold = True
    p.add_run(dict1['c19'])

    p = document.add_paragraph('')
    p.add_run('Work information - Position/Job title: ').bold = True
    p.add_run(dict1['c20'])

    p = document.add_paragraph('')
    p.add_run('Work information - Key job responsibilities: ').bold = True
    p.add_run(dict1['c21'])

    p = document.add_paragraph('')
    p.add_run('Work information - How many years have you been in this position? ').bold = True
    p.add_run(dict1['c22'])

    p = document.add_paragraph('')
    p.add_run('Work information - Work Address: ').bold = True
    p.add_run(dict1['c23'])

    p = document.add_paragraph('')
    p.add_run('Work information - Work Phone Number: ').bold = True
    p.add_run(dict1['c24'])

    p = document.add_paragraph('')
    p.add_run('Work information - Work Email Address: ').bold = True
    p.add_run(dict1['c25'])

    p = document.add_paragraph('')
    p.add_run('List the previous two positions you held before this one:(Please mention name of organization, your position and job responsibilities)').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c26'])

    p = document.add_paragraph('')
    p.add_run('Choose the category that best describes your organization: ').bold = True
    p.add_run(dict1['c27'])

    p = document.add_paragraph('')
    p.add_run('Choose the category that best describes your organization\'s location: ').bold = True
    p.add_run(dict1['c28'])

    p = document.add_paragraph('')
    p.add_run('If you/your company/organization have/has a website, blog, Twitter account or a Facebook page, provide those links here: ').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c29'])

    p = document.add_paragraph('')
    p.add_run('Tell us about your Most Recent Academic experience: - Name of University/College: ').bold = True
    p.add_run(dict1['c30'])

    p = document.add_paragraph('')
    p.add_run('Have you ever traveled internationally? ').bold = True
    p.add_run(dict1['c31'])

    p = document.add_paragraph('')
    p.add_run('Provide details for countries traveled to in the last three years: ').bold = True
    p.add_run(dict1['c32'])

    p = document.add_paragraph('')
    p.add_run('Have you traveled to the U.S. before? ').bold = True
    p.add_run(dict1['c33'])

    p = document.add_paragraph('')
    p.add_run('If yes, please provide information regarding your past travel to the United States.  This should include any travel to the U.S. for school, training, business, or personal reasons.  Please provide dates, reason for travel, and the type of visa you traveled on.  If you have travel planned, but have not yet traveled, please provide that information as well. Please provide the source of funding for your trip.  For example, if you traveled to the U.S. for school, who funded your schooling?').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c34'])

    p = document.add_paragraph('')
    p.add_run('Describe your professional goals in detail: ').bold = True
    p.add_run(dict1['c35'])

    p = document.add_paragraph('')
    p.add_run('Give examples from your work, volunteer, educational or travel experiences that describe in detail your ability to lead, think innovatively, adapt to new situations and work effectively in teams.').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c36'])

    p = document.add_paragraph('')
    p.add_run('In your opinion, what are the most pressing challenges facing economic development and empowerment in your country? \nExplain how you see a role for yourself as an emerging leader in either government, civil society or the private sector in addressing those challenges.').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c37'])

    p = document.add_paragraph('')
    p.add_run('In your opinion, what are the most exciting opportunities facing economic development and empowerment in your country? How do you see yourself contributing to those opportunities being realized?').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c38'])

    p = document.add_paragraph('')
    p.add_run('An important element of being selected for, participating in and continuing future contact with the grantee organization, its partner organizations, the U.S. State Department, other program participants and host organizations is the development, implementation and continuation of an individual project idea. This project should be a natural extension/progression of your current business/organizational/program focus that you believe will be enhanced/improved through your participation in the professional exchanges program. \nUse between 500 and 800 words and describe the nature, goals/objectives and other elements of the project you would like to develop and implement should you be selected for the program. Focus on specific aspects of the project and also discuss how you see a potential host organization in the U.S. (similar to the organization/company you currently work in) being involved with and contributing to the development and implementation of such a project. Also highlight how you will use the different networks you develop while on the professional exchanges program to successfully implement and continue this project. Finally, outline how you would involve people from the local community in your project, how they would benefit and how the U.S. networks could contribute toward benefiting the local communities through your project.').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c39'])

    p = document.add_paragraph('')
    p.add_run('Please write a short biographic paragraph, written in the third person, that would be used in program materials to tell others about you, should you be chosen for the program. (50-100 words):').bold = True
    p = document.add_paragraph('')
    p.add_run(dict1['c40'])

    document.save('data/'+dict1['c1']+'.docx')